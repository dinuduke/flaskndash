import pandas as pd
import numpy as np
import docx
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer


UPLOAD_FOLDER = 'static/uploads'

def frequency_sim():
    doc = docx.Document('static/uploads/Human_resources.docx')
    doc1 = [p.text for p in doc.paragraphs if p.text] 
    x1 = ' '.join(map(str, doc1)) 
    doc_1 = ' '.join(x1.split())

    # print(doc1)

    doc_3 = docx.Document('static/uploads/Python_wiki.docx')
    doc3 = [p.text for p in doc_3.paragraphs if p.text] 
    x3 = ' '.join(map(str, doc3)) 
    doc_3 = ' '.join(x3.split())
    # print(doc_3)

    d = {}
    d[1]= (doc_1)
    d[2]= (doc_3)
    df = pd.DataFrame(d,[0]).T
    df.columns = ["docs"]
    df['docs'] = df['docs'].str.lower()

    ted = df['docs']

    # Create TfidfVectorizer object
    vectorizer = TfidfVectorizer()

    # Generate matrix of word vectors
    tfidf_matrix = vectorizer.fit_transform(ted)

    # Print the shape of tfidf_matrix
    # print(tfidf_matrix.shape)
    x = cosine_similarity(tfidf_matrix)
    data_cor = pd.DataFrame(x)
    data_cor = data_cor*5
    data_cor = data_cor.round(2)
    data_cor = data_cor.mask(data_cor < 0, 0)
    data_cor.columns = ['HR','python']
    data_cor.index = ['HR','python']

    # print(data_cor)

    return data_cor


# print(frequency_sim())